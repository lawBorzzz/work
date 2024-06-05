import os
import pandas as pd
import configparser
import tkinter as tk
from tkinter import ttk, Scrollbar, filedialog, messagebox
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION

def split_text(text, max_length=50):
    words = text.split()
    lines = []
    current_line = ""
    for word in words:
        if len(current_line) + len(word) + 1 <= max_length:
            current_line += " " + word
        else:
            lines.append(current_line.strip())
            current_line = word
    lines.append(current_line.strip())
    return "\n".join(lines)

class MailRegistryApp:
    def __init__(self, master):
        self.master = master
        self.master.title("Менеджер писем")
        self.master.geometry("600x270")  # Устанавливаем размер окна
        self.center_window(self.master)  # Центрируем окно по горизонтали и вертикали

        # Настройка стилей
        self.style = ttk.Style()
        self.style.configure('TButton', font=('Arial', 12), padding=10)
        self.style.configure('TLabel', font=('Arial', 12))

        # Путь к файлу базы данных
        self.registry_path = ""

        # Создаем основной фрейм
        main_frame = ttk.Frame(master)
        main_frame.pack(expand=True)

        # Заголовок
        #title = ttk.Label(main_frame, text="Менеджер писем", font=('Arial', 18, 'bold'))
        #title.pack(pady=20)

        # Фрейм для кнопок
        button_frame = ttk.Frame(main_frame)
        button_frame.pack(pady=10)

        # Кнопки
        self.btn_create_registry = ttk.Button(button_frame, text="Создать список писем", command=self.create_registry)
        self.btn_create_registry.pack(pady=10, ipadx=10)

        self.btn_settings = ttk.Button(button_frame, text="Настройки", command=self.open_settings)
        self.btn_settings.pack(pady=10, ipadx=10)

        # Надпись "by.Borzzz" в нижнем правом углу
        by_label = ttk.Label(master, text="by.Borzzz", foreground="gray")
        by_label.pack(side=tk.RIGHT, padx=10, pady=10)

        # Надпись "Для Грищенко" в нижнем правом углу
        by_label = ttk.Label(master, text="Разработано для Грищенко Е.В.", foreground="gray")
        by_label.pack(side=tk.LEFT, padx=10, pady=10)

        # Инициализация файла конфигурации
        self.config = configparser.ConfigParser()
        self.config_file = os.path.join(os.getenv("APPDATA"), "MailRegistryApp", "config.ini")
        if os.path.exists(self.config_file):
            self.config.read(self.config_file)
            if 'Paths' in self.config:
                self.registry_path = self.config['Paths'].get('registry', '')

    def center_window(self, window):
        # Получаем размеры экрана
        screen_width = window.winfo_screenwidth()
        screen_height = window.winfo_screenheight()

        # Получаем размеры окна
        window.update_idletasks()  # Обновляем информацию о размере окна
        window_width = window.winfo_width()
        window_height = window.winfo_height()

        # Вычисляем координаты для размещения окна по центру экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # Устанавливаем положение окна
        window.geometry(f"{window_width}x{window_height}+{x}+{y}")

    def create_registry(self):
        if not self.registry_path:
            print("Сначала выберите файл базы данных реестра через настройки")
            messagebox.showwarning("Внимание", "Сначала выберите файл базы данных реестра через настройки.")
            return

        # Чтение данных из файла
        try:
            if self.registry_path.endswith('.xls'):
                df = pd.read_excel(self.registry_path, engine='xlrd')
            else:
                df = pd.read_excel(self.registry_path)  # Поддерживаем различные форматы Excel
            
            adresat_column = df['ADRESAT']  # Выбираем столбец с именами
            adresat_values = adresat_column.tolist()  # Преобразуем в список

            # Открываем окно для выбора имен
            self.select_names_window(adresat_values)
        except Exception as e:
            print("Ошибка чтения файла реестра писем:", e)
            messagebox.showerror("Ошибка", f"Ошибка чтения файла реестра писем: {e}")

    def select_names_window(self, adresat_values):
        # Создаем новое окно для выбора имен
        select_names_window = tk.Toplevel(self.master)
        select_names_window.title("Выбор имен")
        select_names_window.geometry("600x800")
        self.center_window(select_names_window)  # Центрируем окно выбора имен

        # Создаем фреймы для структурирования интерфейса
        frame_top = ttk.Frame(select_names_window, padding="10")
        frame_top.pack(fill=tk.BOTH, expand=True)

        frame_bottom = ttk.Frame(select_names_window, padding="10")
        frame_bottom.pack(fill=tk.BOTH, expand=True)

        # Заголовок окна выбора имен
        label = ttk.Label(frame_top, text="Выберите имена из списка", font=('Arial', 14, 'bold'))
        label.pack(pady=10)

        # Сортируем имена в алфавитном порядке
        adresat_values.sort()

        # Создаем список имен с прокруткой
        lb_names = tk.Listbox(frame_top, selectmode=tk.MULTIPLE, width=50, height=20)
        for name in adresat_values:
            lb_names.insert(tk.END, name)
        lb_names.pack(pady=10, side=tk.LEFT, fill=tk.BOTH, expand=True)

        # Создаем и настраиваем скроллбар для списка имен
        scrollbar_names = Scrollbar(frame_top, orient=tk.VERTICAL, command=lb_names.yview)
        scrollbar_names.pack(side=tk.LEFT, fill=tk.Y)
        lb_names.config(yscrollcommand=scrollbar_names.set)

        # Создаем окно для отображения выбранных имен с прокруткой
        selected_names_listbox = tk.Listbox(frame_bottom, width=50, height=10, state='disabled')
        selected_names_listbox.pack(pady=10, side=tk.LEFT, fill=tk.BOTH, expand=True)

        scrollbar_selected = Scrollbar(frame_bottom, orient=tk.VERTICAL, command=selected_names_listbox.yview)
        scrollbar_selected.pack(side=tk.LEFT, fill=tk.Y)
        selected_names_listbox.config(yscrollcommand=scrollbar_selected.set)

        # Счетчик количества выбранных имен
        self.selected_count_var = tk.StringVar()
        self.selected_count_var.set("Выбрано: 0")
        label_selected_count = ttk.Label(frame_bottom, textvariable=self.selected_count_var, font=('Arial', 12))
        label_selected_count.pack(pady=5)

        # Кнопка для создания списка писем
        btn_create_list = ttk.Button(frame_bottom, text="Создать список писем", command=lambda: self.create_word_document (lb_names))
        btn_create_list.pack(pady=5, ipadx=10)

        # Кнопка для очистки списка
        btn_clear_list = ttk.Button(frame_bottom, text="Очистить список", command=lambda: self.clear_selected_names(lb_names, selected_names_listbox))
        btn_clear_list.pack(pady=5, ipadx=10)

        # Кнопка для выхода из программы
        btn_exit = ttk.Button(frame_bottom, text="Выход", command=self.master.quit)
        btn_exit.pack(pady=5, ipadx=10)

        # Обновление окна выбранных имен при изменении выбора в lb_names
        def update_selected_names_listbox(event):
            selected_names_listbox.config(state='normal')
            selected_names_listbox.delete(0, tk.END)
            selected_indices = lb_names.curselection()
            for idx in selected_indices:
                selected_names_listbox.insert(tk.END, lb_names.get(idx))
            selected_names_listbox.config(state='disabled')
            self.selected_count_var.set(f"Выбрано: {len(selected_indices)}")

        lb_names.bind('<<ListboxSelect>>', update_selected_names_listbox)

    def clear_selected_names(self, lb_names, selected_names_listbox):
        lb_names.selection_clear(0, tk.END)
        selected_names_listbox.delete(0, tk.END)
        self.selected_count_var.set("Выбрано: 0")

    def create_word_document(self, lb_names):
        # Получаем выбранные имена
        selected_indices = lb_names.curselection()
        selected_names = [lb_names.get(idx) for idx in selected_indices]
        print("Выбранные имена:", selected_names)

        if not selected_names:
            print("Выберите хотя бы одно имя из списка")
            return

        # Чтение данных из файла
        try:
            df = pd.read_excel(self.registry_path)  # Поддерживаем различные форматы Excel

            # Создаем новый документ Word
            # Создаем новый документ Word
            doc = Document()

            # Устанавливаем ориентацию и размеры страницы
            sections = doc.sections
            for section in sections:
                section.orientation = WD_ORIENTATION.LANDSCAPE
                section.page_width = Cm(22.9)  # Устанавливаем ширину листа
                section.page_height = Cm(16.2)  # Устанавливаем высоту листа
                section.left_margin = Cm(1)  # Устанавливаем левый отступ
                section.right_margin = Cm(1)  # Устанавливаем правый отступ
                section.top_margin = Cm(1)  # Устанавливаем верхний отступ
                section.bottom_margin = Cm(1)  # Устанавливаем нижний отступ
            
            # Добавляем в верхний левый угол на всех листах информацию об отправителе
            sender_info = " От кого: Девятнадцатый арбитражный\n                           апелляционный суд\n Откуда: г. Воронеж, ул. Платонова, д. 8"
            for section in doc.sections:
                header = section.header
                paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT  # Выравниваем по центру
                paragraph.text = sender_info

            for name in selected_names:
                # Фильтруем строки по выбранному имени
                selected_rows = df[df['ADRESAT'] == name]

                # Создаем новый параграф
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю

                # Добавляем пропуски перед данными
                for _ in range(14):
                    p.add_run("\n")

                # Добавляем данные из выбранных строк
                for index, row in selected_rows.iterrows():
                    # Добавляем адресат_2
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю
                        value = row['ADRESAT_2']
                        if not pd.isnull(value):
                            # Разбиваем текст на строки, если больше 75 символов
                            if len(value) > 75:
                                value = split_text(value)
                            run = p.add_run(f"{value}")
                            run.font.size = Pt(10)  # Устанавливаем размер шрифта
                            run.font.italic = True  # Устанавливаем курсив
                        
                    # Добавляем адресат
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю
                        value = row['ADRESAT']
                        if not pd.isnull(value):
                            run = p.add_run(f"{value}")
                            run.font.size = Pt(10)  # Устанавливаем размер шрифта
                            run.font.bold = True  # Устанавливаем жирный шрифт
                            run.font.italic = True  # Устанавливаем курсив

                        # Добавляем адрес
                        p = doc.add_paragraph()
                        p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT  # Выравниваем по правому краю
                        value = row['ADDRESSLINE']
                        if not pd.isnull(value):
                            run = p.add_run(f"{value}")
                            run.font.size = Pt(10)  # Устанавливаем размер шрифта
                            run.font.italic = True  # Устанавливаем курсив

                # Добавляем разрыв страницы между именами
                if name != selected_names[-1]:
                    doc.add_page_break()

            # Сохраняем документ
            output_path = filedialog.asksaveasfilename(defaultextension=".docx")
            if output_path:
                doc.save(output_path)
                print("Документ успешно создан и сохранен в файле:", output_path)
        except Exception as e:
            print("Ошибка при создании документа Word:", e)

    def open_settings(self):
        # Создаем новое окно настроек
        settings_window = tk.Toplevel(self.master)
        settings_window.title("Настройки")
        settings_window.geometry("400x250")

        # Поля для путей к базам данных
        registry_path_var = tk.StringVar()

        # Загрузка путей из конфигурации
        if 'Paths' in self.config:
            registry_path_var.set(self.config['Paths'].get('registry', ''))

        # Функции выбора пути для баз данных
        def choose_registry_path():
            path = filedialog.askopenfilename(
                title="Выберите файл базы данных реестра",
                filetypes=[("Excel files", "*.xls;*.xlsx"), ("All files", "*.*")])

            if path:
                registry_path_var.set(path)
                save_paths()  # Сохраняем путь при выборе файла

        # Функция сохранения путей в файл конфигурации
        def save_paths():
            self.registry_path = registry_path_var.get()  # Обновляем путь к файлу реестра
            self.config['Paths'] = {
                'registry': registry_path_var.get(),
            }
            with open(self.config_file, 'w') as configfile:
                self.config.write(configfile)
            settings_window.destroy()  # Закрываем окно настроек после сохранения

        # Виджеты настроек
        frame_settings = ttk.Frame(settings_window, padding="10")
        frame_settings.pack(fill=tk.BOTH, expand=True)

        label_registry_path = ttk.Label(frame_settings, text="Путь к файлу БД:", font=('Arial', 12))
        label_registry_path.pack(pady=5)

        entry_registry_path = ttk.Entry(frame_settings, textvariable=registry_path_var, state='readonly', width=50)
        entry_registry_path.pack(pady=5)

        btn_registry_path = ttk.Button(frame_settings, text="Выбрать файл", command=choose_registry_path)
        btn_registry_path.pack(pady=5, ipadx=10)

        btn_save = ttk.Button(frame_settings, text="Сохранить", command=save_paths)
        btn_save.pack(pady=10, ipadx=10)

        by_label = ttk.Label(frame_settings, text="Версия: 1.3", foreground="gray")
        by_label.pack(side=tk.RIGHT, padx=10, pady=10)

        # Устанавливаем начальные значения в полях
        entry_registry_path.insert(0, registry_path_var.get())

        # Устанавливаем окно настроек модальным
        settings_window.transient(self.master)
        settings_window.grab_set()
        self.master.wait_window(settings_window)

def main():
    root = tk.Tk()
    app = MailRegistryApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
