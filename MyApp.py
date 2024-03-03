import calendar
import io
import json
import math
import os
import re
import requests
import subprocess
import tempfile
import webbrowser
import tkinter as tk

from collections import defaultdict
from datetime import date, datetime
from tkinter import Tk, Button, messagebox, simpledialog
from tkinter.filedialog import askdirectory
from tkinter import PhotoImage
from tkinter import simpledialog
from tkinter import font
from PIL import Image, ImageTk
from tkcalendar import Calendar, DateEntry
from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Cm, Pt, RGBColor


def check_for_updates(current_version):
    try:
        response = requests.get('https://raw.githubusercontent.com/lawBorzzz/work/main/version.txt')
        response.raise_for_status()
        latest_version = response.text.strip()
        print("Получена версия из GitHub:", latest_version)
    except requests.exceptions.RequestException as e:
        print("Ошибка при проверке обновлений:", e)
        return

    if latest_version != current_version:
        print("Доступна новая версия:", latest_version)
        # Открываем диалоговое окно с предложением обновления
        root = tk.Tk()
        root.withdraw()  # Скрыть основное окно
        response = messagebox.askquestion("Доступна новая версия", "Желаете обновить?")
        if response == 'yes':
            # Перенаправляем пользователя на страницу загрузки
            webbrowser.open('https://github.com/lawBorzzz/work/releases')
        root.destroy()
    else:
        print("У вас последняя версия приложения.")

class App(tk.Tk):
    BASE_COST = 89.5  # базовая стоимость бандероли
    STEP_COST = 3.5   # стоимость за шаг в 20 грамм
    LETTER_COST = 29.0  # стоимость письма простого
    REGISTERED_LETTER_COST = 67.0 # стоимость письма заказного
    NDS = 1.2 # НДС 20%


    
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self.title("Формирование отчета")

        # Получить путь к исполняемому файлу
        executable_path = os.path.dirname(os.path.abspath(__file__))

        # Путь к файлу background.jpg в папке с исполняемым файлом
        background_image_path = os.path.join(executable_path, 'background.jpg')

        # Создаем объект Image для фона
        background_image = Image.open(background_image_path)
        background_image = background_image.resize((900, 700), Image.LANCZOS)

        # Создаем объект PhotoImage для фона
        self.bg_image_tk = ImageTk.PhotoImage(background_image)

        self.load_settings_from_file() # загружаем настройки при старте приложения

        self.total_weight = 0
        self.total_cost = 0.0
        self.total_parcels = 0

        self.weights = []

        self.numbers_entered = []  # Список для хранения введенных значений простых писем
        self.numbers_entered_reg = []  # Список для хранения введенных значений заказных писем

        # Получить путь к рабочему столу пользователя
        self.desktop_path = os.path.join(os.path.expanduser('~'), 'Desktop')
        
        self.create_widgets()

    # Центрирование окна
    def center_window(self, width, height):
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        x = (screen_width // 2) - (width // 2)
        y = (screen_height // 2) - (height // 2)

        self.geometry(f"{width}x{height}+{x}+{y}")

    # Создание виджетов
    def create_widgets(self):
        self.geometry("600x305")  # Задаем размер корневого окна

        # Получить путь к исполняемому файлу
        executable_path = os.path.dirname(os.path.abspath(__file__))

        # Путь к файлу settings_image.png в папке с исполняемым файлом
        settings_image_path = os.path.join(executable_path, 'settings_image.png')

        # Создаем метку с изображением в качестве фона
        self.background_label = tk.Label(self.master, image=self.bg_image_tk)
        self.background_label.place(x=0, y=0, relwidth=1, relheight=1)

        button_frame = tk.Frame(self)
        button_frame.pack(pady=10)

        # Кнопки
        button_texts = ["Подсчет бандеролей", "     Подсчет писем     ", "   Подсчет посылок   ", "      Отчет за месяц     ", "    Создать обложку   "]
        button_commands = [self.open_packet_window, self.open_letters_window, self.open_parcels_window,
                           self.ask_month_input, self.open_cover_window]

        for text, command in zip(button_texts, button_commands):
            button = tk.Button(self, text=text, command=command)
            button.pack(pady=10)
            button.configure(borderwidth=2, relief=tk.GROOVE, bg='lightgrey')

        # Создаем кнопку настроек с изображением
        settings_image = Image.open(settings_image_path)
        settings_image_tk = ImageTk.PhotoImage(settings_image)

        self.settings_button = tk.Button(self, image=settings_image_tk, command=self.open_settings_window)
        self.settings_button.image = settings_image_tk  # сохраняем ссылку на объект
        self.settings_button.pack(side=tk.BOTTOM, anchor=tk.SW, padx=11, pady=10)
        self.settings_button.configure(borderwidth=2, relief=tk.GROOVE, bg='lightgrey')

        self.center_window(600, 305)

# Открывается кнопка ввода веса бандеролей       
    def open_packet_window(self):
        if hasattr(self, 'packet_window') and self.packet_window.winfo_exists():
            self.packet_window.focus_set()
            return
        
        self.packet_window = tk.Toplevel(self)
        self.packet_window.title("Подсчет бандеролей")
        self.packet_window.geometry("300x700")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.packet_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        #Выравнивание посередине
        screen_width = self.packet_window.winfo_screenwidth()
        screen_height = self.packet_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 - 500
        y_coordinate = (screen_height - 700) // 2

        # Устанавливаем положение окна по центру
        self.packet_window.geometry(f"300x700+{x_coordinate}+{y_coordinate}")

        self.packet_label = tk.Label(self.packet_window, text="Введите вес бандероли (в граммах):")
        self.packet_label.pack(pady=10)

        self.packet_entry = tk.Entry(self.packet_window)
        self.packet_entry.pack(pady=10, padx=20)
        self.packet_entry.focus_set()

        # Создание Listbox для отображения введенных весов
        self.packets_listbox_label = tk.Label(self.packet_window, text="Список введённых бандеролей:")
        self.packets_listbox_label.pack(pady=(20, 5), padx=10, anchor=tk.W)
        self.packets_listbox = tk.Listbox(self.packet_window)
        self.packets_listbox.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        self.delete_selected_weight_button = tk.Button(
            self.packet_window, text="Удалить выбранный вес", command=self.delete_selected_weight, bg='lightgrey'
        )
        self.delete_selected_weight_button.pack(pady=10, padx=20, fill=tk.X)
        self.delete_selected_weight_button.configure(borderwidth=2, relief=tk.GROOVE)

        self.finish_button = tk.Button(
            self.packet_window, text="Завершить подсчет", command=self.finish_weight_calculation, bg='lightgrey'
        )
        self.finish_button.pack(pady=10, padx=20, fill=tk.X)
        self.finish_button.configure(borderwidth=2, relief=tk.GROOVE)

        self.packet_entry.bind("<Return>", self.add_weight)

# Это округление введенных бандеролей до целого четного числа равному 20
    def round_weight(self, weight):
        return math.ceil(weight / 20.0) * 20

# Добавление бандеролей в общий список до подсчета
    def add_weight(self, event=None):
        try:
            weight_str = self.packet_entry.get().strip()
            if not weight_str:
                raise ValueError("Введите корректное числовое значение.")
            
            if not re.match(r'^\d*\.?\d*$', weight_str):
                raise ValueError("Введите только цифры!")
            
            weight = float(weight_str)
            if weight < 120 or weight > 2000:
                raise ValueError("Введите валидный вес (от 120 до 2000).")
        
            # Округляем вес
            rounded_weight = self.round_weight(weight)
            self.weights.append(rounded_weight)  # Добавляем округленный вес
            self.total_weight += rounded_weight  # Используем округленный вес для общего веса
            self.total_parcels += 1
            self.total_cost += self.calculate_cost(rounded_weight)  # Рассчитываем стоимость по округленному весу
            self.packets_listbox.insert(tk.END, f"{rounded_weight} грамм")
            self.packet_entry.delete(0, tk.END)
        except ValueError as e:
            messagebox.showwarning("Ошибка", str(e))
        finally:
            self.packet_entry.focus_set()

# Удаление выбранного веса из списка   
    def delete_selected_weight(self):
        selection = self.packets_listbox.curselection()  # Получаем текущий выбранный элемент в listbox
        if selection:
            index = selection[0]
            weight_to_remove = self.weights.pop(index)  # Удалить вес из списка
            self.total_weight -= weight_to_remove
            self.total_cost -= self.calculate_cost(weight_to_remove)
            self.packets_listbox.delete(index)  # Удалить элемент из listbox
            self.total_parcels -= 1
        else:
            messagebox.showwarning("Ошибка", "Выберите значение для удаления.")
        # Фокусировка после изменения
        self.packet_entry.focus_set()

    def finish_weight_calculation(self):
        self.open_date_window()

# Расчет стоимости бандеролей, идет по настройкам с выставленными значениями.
    def calculate_cost(self, weight):
        additional_cost = max(0, (weight - 120) // 20 * self.STEP_COST)
        return self.BASE_COST + additional_cost

# Открывается поле ввода даты перед сохранением списка с бандеролями       
    def open_date_window(self):
        if hasattr(self, 'date_window') and self.date_window.winfo_exists():
            self.date_window.focus_set()
            return
        self.date_window = tk.Toplevel(self)
        self.date_window.title("Выберите дату")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.date_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)
       
        window_width = 300
        window_height = 280

        # Получаем размеры экрана
        screen_width = self.date_window.winfo_screenwidth()
        screen_height = self.date_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # Устанавливаем положение окна
        self.date_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Создаем календарь
        self.cal = Calendar(self.date_window, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Функция сохранения результатов при выборе даты
        def save_results():
            selected_date = self.cal.get_date()
            self.save_results(selected_date)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(
        self.date_window, text="Сформировать список", command=save_results, bg='lightgray')
        self.save_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_button.configure(borderwidth=2, relief=tk.GROOVE)

# Сохранение результата списка бандеролей с датой
    def save_results(self, event=None):
        if not self.weights:
            messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.packet_entry.focus_set()
            return
    
        selected_date = self.cal.get_date()
        try:
            current_date = datetime.strptime(selected_date, "%d.%m.%Y").date()
            
            result_string = (f"Итого за {current_date.strftime('%d.%m.%Y')} отправлено {self.total_parcels}"
                 f" {'бандероль' if self.total_parcels % 10 == 1 and self.total_parcels % 100 != 11 else 'бандероли' if 2 <= self.total_parcels % 10 <= 4 and (self.total_parcels % 100 < 10 or self.total_parcels % 100 >= 20) else 'бандеролей'} весом {self.total_weight:.2f} грамм на сумму {self.total_cost:.2f} руб.\n")
            
            custom_path = self.custom_path
            filename = os.path.join(custom_path, f"Списки бандеролей.txt")

            with open(filename, "a", encoding='utf-8') as file:
                file.write(result_string)
        
            messagebox.showinfo("Успешно", "Результаты сохранены.")
            self.date_window.destroy()
            self.packet_window.destroy()
        
            self.weights.clear()
            self.total_weight = 0
            self.total_cost = 0.0
            self.total_parcels = 0
        except ValueError:
            messagebox.showerror("Ошибка", "Введите дату в правильном формате (дд.мм.гггг).")

# Окно выбора типа письма.
    def open_letters_window(self):
        if hasattr(self, 'letters_window') and self.letters_window.winfo_exists():
            self.letters_window.focus_set()
            return

        self.letters_window = tk.Toplevel(self)
        self.letters_window.title("Выберите тип письма")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.letters_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Размер окна
        window_width = 400
        window_height = 200

        # Получаем размеры экрана
        screen_width = self.letters_window.winfo_screenwidth()
        screen_height = self.letters_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2 - 285

        # Устанавливаем позицию окна
        self.letters_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Кнопки
        self.simple_button = tk.Button(self.letters_window, text="   Простые письма   ", command=self.calculate_simple_letters)
        self.simple_button.pack(pady=(20, 0))  # Устанавливаем отступ сверху для первой кнопки
        self.simple_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.simple_button.configure(bg='lightgrey')

        tk.Label(self.letters_window, text="").pack()

        self.registered_button = tk.Button(self.letters_window, text="  Заказные письма   ", command=self.calculate_registered_letters)
        self.registered_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.registered_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.registered_button.configure(bg='lightgrey')

        tk.Label(self.letters_window, text="").pack()

        self.foreign_button = tk.Button(self.letters_window, text="      Иностранные      ", command=self.calculate_foreign_letters)
        self.foreign_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.foreign_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.foreign_button.configure(bg='lightgrey')

        self.letters_window.focus_set()

# Открывается подсчет иностранных писем
    def calculate_foreign_letters(self):
        if hasattr(self, 'foreign_window') and self.foreign_window.winfo_exists():
            self.foreign_window.focus_set()
            return

        # Открытие нового окна для ввода цены и даты
        self.foreign_window = tk.Toplevel(self)
        self.foreign_window.title("Подсчет иностранных писем")
        self.foreign_window.geometry("300x700")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.foreign_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.foreign_window.winfo_screenwidth()
        screen_height = self.foreign_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 - 500
        y_coordinate = (screen_height - 700) // 2

        # Устанавливаем положение окна по центру
        self.foreign_window.geometry(f"300x700+{x_coordinate}+{y_coordinate}")

        self.prices_entered = []  # Список для хранения введенных цен

        # Ввод цены
        self.price_label = tk.Label(self.foreign_window, text="Введите цену письма:")
        self.price_label.pack(pady=(20, 5))

        self.price_entry = tk.Entry(self.foreign_window)
        self.price_entry.pack(pady=5)
        self.price_entry.bind("<Return>", self.add_to_foreign_list)
        self.price_entry.focus_set()

        # Список введенных цен
        self.listbox_label = tk.Label(self.foreign_window, text="Список введенных цен:")
        self.listbox_label.pack(pady=(20, 5), padx=10, anchor=tk.W)
        self.listbox = tk.Listbox(self.foreign_window)
        self.listbox.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        self.delete_selected_button = tk.Button(
            self.foreign_window, text="Удалить выбранное", command=self.remove_foreign_selected, bg='lightgrey'
        )
        self.delete_selected_button.pack(pady=5, padx=20, fill=tk.X)
        self.delete_selected_button.configure(borderwidth=2, relief=tk.GROOVE)

        self.finish_button = tk.Button(
            self.foreign_window, text="Завершить подсчет", command=self.open_foreign_calendar, bg='lightgrey'
        )
        self.finish_button.pack(pady=5, padx=20, fill=tk.X)
        self.finish_button.configure(borderwidth=2, relief=tk.GROOVE)

# добавляем в листбокс
    def add_to_foreign_list(self, event=None):
        try:
            price = self.price_entry.get().replace(',', '.')  # Заменяем запятую на точку
            if not price:
                raise ValueError("Пожалуйста, введите цену.")
            self.price_entry.focus_set()

            # Преобразуем строку в число
            price_float = float(price)

            if price_float <= 0:
                raise ValueError("Введите положительное числовое значение.")
            self.price_entry.focus_set()

            # Проверяем, что ввод содержит только цифры и точку
            if not re.match(r'^\d*\.?\d*$', price):
                raise ValueError("Введите корректное числовое значение.")
            self.price_entry.focus_set()

            self.prices_entered.append(price_float)
            self.listbox.insert(tk.END, f"{price_float} руб.")
            self.price_entry.delete(0, tk.END)
        except ValueError as e:
            messagebox.showerror("Ошибка", "Введите корректное число!")
        finally:
            self.price_entry.focus_set()

# Удаляем из листбокс
    def remove_foreign_selected(self):
        selected_indices = self.listbox.curselection()
        if selected_indices:
            selected_index = selected_indices[0]
            self.listbox.delete(selected_index)
            del self.prices_entered[selected_index]
        else:
            tk.messagebox.showwarning("Предупреждение", "Выберите значение для удаления.")
            self.price_entry.focus_set()

# Открываем календарь
    def open_foreign_calendar(self):
        if hasattr(self, 'calendar_window') and self.calendar_window.winfo_exists():
            self.calendar_window.focus_set()
            return
        self.calendar_window = tk.Toplevel(self.foreign_window)
        self.calendar_window.title("Выберите дату")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.calendar_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        window_width = 300
        window_height = 280

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2
        y = (screen_height - window_height) // 2

        # Устанавливаем положение окна
        self.calendar_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Создание календаря
        self.calendar = Calendar(self.calendar_window, locale='ru_RU')
        self.calendar.pack(padx=20, pady=20)

        # Кнопка "Сохранить дату"
        self.save_button = tk.Button(self.calendar_window, text="Сформировать список", command=self.save_foreign_date, bg='lightgray')
        self.save_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_button.configure(borderwidth=2, relief=tk.GROOVE)

# Сохраняем результат
    def save_foreign_date(self):
        if not self.prices_entered:
            messagebox.showerror("Ошибка", "Список введенных цен пуст.")
            self.price_entry.focus_set()
            return
        
    
        selected_date = self.calendar.get_date()
        self.save_to_foreign_file(selected_date)
        self.calendar_window.destroy()
        messagebox.showinfo("Успешно", "Данные сохранены в файл.")

        # Закрываем окно
        self.foreign_window.destroy()
        self.letters_window.focus_set()

# Сохраняем в файл
    def save_to_foreign_file(self, date):
        filename = os.path.join(self.custom_path, "Списки иностранных писем.txt")
        with open(filename, "a", encoding='utf-8') as file:
            for price in self.prices_entered:
                file.write(f"За {date}: 1 письмо, ценой: {price} руб.\n")

# Отыкрывается кнопка ЗАКАЗНЫХ писем  
    def calculate_registered_letters(self):
        if hasattr(self, 'registered_window') and self.registered_window.winfo_exists():
            self.registered_window.focus_set()
            return

        self.registered_window = tk.Toplevel()
        self.registered_window.title("Подсчет заказных писем")
        self.registered_window.geometry("300x700")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.registered_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.registered_window.winfo_screenwidth()
        screen_height = self.registered_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 - 500
        y_coordinate = (screen_height - 700) // 2

        # Устанавливаем положение окна по центру
        self.registered_window.geometry(f"300x700+{x_coordinate}+{y_coordinate}")

        self.numbers_entered_reg = []  # Список для хранения введенных значений

        # Ввод количества писем
        self.quantity_label = tk.Label(self.registered_window, text="Введите количество писем:")
        self.quantity_label.pack(pady=(20, 5))

        self.quantity_entry = tk.Entry(self.registered_window)
        self.quantity_entry.pack(pady=5)
        self.quantity_entry.bind("<Return>", self.add_to_list_reg)  # Привязка к кнопке Enter
        self.quantity_entry.focus_set()

        # Список введенных значений
        self.listbox_label = tk.Label(self.registered_window, text="Список введенных писем:")
        self.listbox_label.pack(pady=(20, 5), padx=10, anchor=tk.W)
        self.listbox = tk.Listbox(self.registered_window)
        self.listbox.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        # Кнопка "удалить выбранное" для удаления конкретного введенного результата
        self.delete_selected_button_reg = tk.Button(self.registered_window, text="Удалить выбранное",
                                                     command=self.remove_selected_reg, bg='lightgrey')
        self.delete_selected_button_reg.pack(pady=5, padx=20, fill=tk.X)
        self.delete_selected_button_reg.configure(borderwidth=2, relief=tk.GROOVE)

        # Кнопка "Закончить подсчет" для открытия календаря
        self.finish_button = tk.Button(self.registered_window, text="Завершить подсчет",
                                        command=self.open_calendar_reg, bg='lightgrey')
        self.finish_button.pack(pady=5, padx=20, fill=tk.X)
        self.finish_button.configure(borderwidth=2, relief=tk.GROOVE)

# Добавление значений в листбокс   
    def add_to_list_reg(self, event):
        # Попытка преобразовать введенные данные в число и добавление в список
        try:
            num_letters = int(self.quantity_entry.get())
            if num_letters <= 0:
                messagebox.showerror("Ошибка", "Введите корректное число!")
                self.quantity_entry.focus_set()
                return
            self.numbers_entered_reg.append(num_letters)  # Добавление числа в список
            self.listbox.insert(tk.END, num_letters)  # Вывод числа в интерфейсе
            self.quantity_entry.delete(0, tk.END)  # Очистка поля ввода
        except ValueError:
            tk.messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.quantity_entry.focus_set()

# Если нажали кнопку удалить последний результат из списка писем.
    def remove_selected_reg(self):
        try:
            # Получить индекс выбранного элемента
            index = self.listbox.curselection()[0]
            # Удалить этот элемент из Listbox и из списка numbers_entered
            self.numbers_entered_reg.pop(index)
            self.listbox.delete(index)
        except IndexError:
            tk.messagebox.showwarning("Предупреждение", "Выберите значение для удаления.")
            self.quantity_entry.focus_set()
        except Exception as e:
            tk.messagebox.showwarning("Ошибка", f"Произошла ошибка: {e}")
            self.quantity_entry.focus_set()

# Создание календаря
    def open_calendar_reg(self):
        if hasattr(self, 'calendar_window') and self.calendar_window.winfo_exists():
            self.calendar_window.focus_set()
            return
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.registered_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x280")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.calendar_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 280) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x280+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сформировать список", command=self.calculate_and_save_result_reg, bg='lightgray')
        self.save_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_button.configure(borderwidth=2, relief=tk.GROOVE)
        
# Подсчет и сохранение итога по письмам
    def calculate_and_save_result_reg(self, event=None):
        # Ввод даты и подсчет итогаa
        selected_date = self.cal.get_date()
        selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
        if not selected_date or not self.numbers_entered_reg:
            messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.quantity_entry.focus_set()
            return

        # Подсчет итога
        total = sum(self.numbers_entered_reg) * self.REGISTERED_LETTER_COST
        self.save_to_file_reg(total, sum(self.numbers_entered_reg), selected_date)

        # Отображение результата
        tk.messagebox.showinfo("Успешно", "Результаты сохранены.")
    
        # Закрытие окна ввода
        self.registered_window.destroy()
        self.letters_window.focus_set()

    def save_to_file_reg(self, total_result, total_registered_letters, date):
        custom_path = self.custom_path
        filename = os.path.join(custom_path, f"Списки заказных писем.txt")

        # Используем режим 'a' для добавления данных в конец файла
        with open(filename, 'a', encoding='utf-8') as file:
            file.write(f"Дата: {date} Количество писем: {total_registered_letters} Итого: {total_result} руб.\n")


# Отыкрывается кнопка ПРОСТЫХ писем  
    def calculate_simple_letters(self):
        if hasattr(self, 'simple_window') and self.simple_window.winfo_exists():
            self.simple_window.focus_set()
            return
   
        self.simple_window = tk.Toplevel()
        self.simple_window.title("Подсчет простых писем")
        self.simple_window.geometry("300x700")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.simple_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.simple_window.winfo_screenwidth()
        screen_height = self.simple_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 - 500
        y_coordinate = (screen_height - 700) // 2

        # Устанавливаем положение окна по центру
        self.simple_window.geometry(f"300x700+{x_coordinate}+{y_coordinate}")

        self.numbers_entered = []  # Список для хранения введенных значений

        # Ввод количества писем
        self.quantity_label = tk.Label(self.simple_window, text="Введите количество писем:")
        self.quantity_label.pack(pady=(20, 5))

        self.quantity_entry = tk.Entry(self.simple_window)
        self.quantity_entry.pack(pady=5)
        self.quantity_entry.bind("<Return>", self.add_to_simple_list)  # Привязка к кнопке Enter
        self.quantity_entry.focus_set()

        # Список введенных значений
        self.listbox_label = tk.Label(self.simple_window, text="Список введённых писем:")
        self.listbox_label.pack(pady=(20, 5), padx=10, anchor=tk.W)
        self.listbox = tk.Listbox(self.simple_window)
        self.listbox.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        # Кнопка "удалить выбранное" для удаления конкретного введенного результата
        self.delete_selected_button = tk.Button(self.simple_window, text="Удалить выбранное",
                                                 command=self.remove_simple_selected, bg='lightgrey')
        self.delete_selected_button.pack(pady=5, padx=20, fill=tk.X)
        self.delete_selected_button.configure(borderwidth=2, relief=tk.GROOVE)

        # Кнопка "Закончить подсчет" для открытия календаря
        self.finish_button = tk.Button(self.simple_window, text="Завершить подсчет",
                                        command=self.open_simple_calendar, bg='lightgrey')
        self.finish_button.pack(pady=5, padx=20, fill=tk.X)
        self.finish_button.configure(borderwidth=2, relief=tk.GROOVE)

# Это лист, где отображаются введенные письма (как в памяти так и в окне в виде списка)    
    def add_to_simple_list(self, event):
        # Попытка преобразовать введенные данные в число и добавление в список
        try:
            num_letters = int(self.quantity_entry.get())
            if num_letters <= 0:
                messagebox.showerror("Ошибка", "Введите корректное число!")
                self.quantity_entry.focus_set()
                return
            self.numbers_entered.append(num_letters)  # Добавление числа в список
            self.listbox.insert(tk.END, num_letters)  # Вывод числа в интерфейсе
            self.quantity_entry.delete(0, tk.END)  # Очистка поля ввода
        except ValueError:
            tk.messagebox.showerror("Ошибки", "Список введенных значений пуст.")
            self.quantity_entry.focus_set()

# Если нажали кнопку удалить последний результат из списка писем.
    def remove_simple_selected(self):
        try:
            # Получить индекс выбранного элемента
            index = self.listbox.curselection()[0]
            # Удалить этот элемент из Listbox и из списка numbers_entered
            self.numbers_entered.pop(index)
            self.listbox.delete(index)
        except IndexError:
            tk.messagebox.showwarning("Предупреждение", "Выберите значение для удаления.")
            self.quantity_entry.focus_set()
        except Exception as e:
            tk.messagebox.showwarning("Ошибка", f"Произошла ошибка: {e}")
            self.quantity_entry.focus_set()

# Создание календаря с выбором даты
    def open_simple_calendar(self):
        if hasattr(self, 'calendar_window') and self.calendar_window.winfo_exists():
            self.calendar_window.focus_set()
            return
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.simple_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x280")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.calendar_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 280) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x280+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year,
                            month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сформировать список", command=self.calculate_and_save_simple_result, bg='lightgrey')
        self.save_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_button.configure(borderwidth=2, relief=tk.GROOVE)

# Подсчет и сохранение итога по письмам
    def calculate_and_save_simple_result(self, event=None):
        # Ввод даты и подсчет итога
        selected_date = self.cal.get_date()
        selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
        if not selected_date or not self.numbers_entered:
            messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.quantity_entry.focus_set()
            return

        try:
            total = sum(self.numbers_entered) * self.LETTER_COST
            self.save_to_simple_file(total, sum(self.numbers_entered), selected_date)

            tk.messagebox.showinfo("Успешно", "Результаты сохранены.")
            
            # Закрываем окно
            self.simple_window.destroy()
            self.letters_window.focus_set()
        except ValueError:
            tk.messagebox.showerror("Ошибка", "Ошибка при сохранении результатов.")

    def save_to_simple_file(self, total_result, total_letters, date):
        custom_path = self.custom_path
        filename = os.path.join(custom_path, f"Списки простых писем.txt")

        # Используем режим 'a' для добавления данных в конец файла
        with open(filename, 'a', encoding='utf-8') as file:
            file.write(f"Дата: {date} Количество писем: {total_letters} Итого: {total_result} руб.\n")

# Подсчет посылок
    def open_parcels_window(self):
        if hasattr(self, 'parcels_window') and self.parcels_window.winfo_exists():
            self.parcels_window.focus_set()
            return

        self.parcels_window = tk.Toplevel(self)
        self.parcels_window.title("Подсчет посылок")
        self.parcels_window.geometry("300x700")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.parcels_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.parcels_window.winfo_screenwidth()
        screen_height = self.parcels_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 - 500
        y_coordinate = (screen_height - 700) // 2

        # Устанавливаем положение окна по центру
        self.parcels_window.geometry(f"300x700+{x_coordinate}+{y_coordinate}")

        self.parcels_price_label = tk.Label(self.parcels_window, text="Введите цену посылки без НДС:")
        self.parcels_price_label.pack(pady=10)

        self.parcels_price_entry = tk.Entry(self.parcels_window)
        self.parcels_price_entry.pack(pady=10)
        self.parcels_price_entry.bind("<Return>", self.add_parcel_weight)
        self.parcels_price_entry.focus_set()

        self.parcels_listbox_label = tk.Label(self.parcels_window, text="Список введённых цен:")
        self.parcels_listbox_label.pack(pady=(20, 5), padx=10, anchor=tk.W)

        # Добавим Frame для отступов вокруг Listbox
        listbox_frame = tk.Frame(self.parcels_window)
        listbox_frame.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        self.parcels_weights_listbox = tk.Listbox(listbox_frame)
        self.parcels_weights_listbox.pack(fill=tk.BOTH, expand=True)

        self.delete_selected_parcel_button = tk.Button(
            self.parcels_window, text="Удалить выбранную посылку", command=self.delete_selected_parcel, bg='lightgrey'
        )
        self.delete_selected_parcel_button.pack(pady=10, padx=20, fill=tk.X)
        self.delete_selected_parcel_button.configure(borderwidth=2, relief=tk.GROOVE)

        
        self.finish_button = tk.Button(self.parcels_window, text="Завершить подсчет",
                                    command=self.open_calendar_parcels, bg='lightgrey')
        self.finish_button.pack(pady=10, padx=20, fill=tk.X)
        self.finish_button.configure(borderwidth=2, relief=tk.GROOVE)

    def add_parcel_weight(self, event=None):
        try:
            # Получаем значение из виджета Entry и заменяем запятую на точку
            price_entry_text = self.parcels_price_entry.get().replace(',', '.')
            price = float(price_entry_text)
            if price <= 0:
                messagebox.showerror("Ошибка", "Введите корректную цену.")
                self.parcels_price_entry.focus_set()
                return
            self.parcels_weights_listbox.insert(tk.END, f"{price} руб.")
            self.parcels_price_entry.delete(0, tk.END)
        except ValueError as e:
            messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.parcels_price_entry.focus_set()
        finally:
            self.parcels_price_entry.focus_set()

    def delete_selected_parcel(self):
        selected_index = self.parcels_weights_listbox.curselection()
        if selected_index:
            self.parcels_weights_listbox.delete(selected_index)
        else:
            messagebox.showwarning("Предупреждение", "Выберите значение для удаления.")
            self.parcels_price_entry.focus_set()

    def open_calendar_parcels(self):
        if hasattr(self, 'calendar_window') and self.calendar_window.winfo_exists():
            self.calendar_window.focus_set()
            return
        # Создание календаря
        self.calendar_window = tk.Toplevel(self.parcels_window)
        self.calendar_window.title("Выберите дату")
        self.calendar_window.geometry("300x280")

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.calendar_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.calendar_window.winfo_screenwidth()
        screen_height = self.calendar_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2
        y_coordinate = (screen_height - 280) // 2

        # Устанавливаем положение окна по центру
        self.calendar_window.geometry(f"300x280+{x_coordinate}+{y_coordinate}")

        self.cal = Calendar(self.calendar_window, selectmode="day", year=datetime.now().year, 
                            month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        self.cal.pack(pady=10)

        # Кнопка для сохранения результатов
        self.save_button = tk.Button(self.calendar_window, text="Сформировать список", command=self.calculate_and_save_parcels, bg='lightgray')
        self.save_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_button.configure(borderwidth=2, relief=tk.GROOVE)

    def calculate_and_save_parcels(self):
        if not self.parcels_weights_listbox.size():
            messagebox.showerror("Ошибка", "Список введенных значений пуст.")
            self.parcels_price_entry.focus_set()
            return
        
        selected_date = self.cal.get_date()
        selected_date = re.sub(r'[\s\\\/.,]', '.', selected_date)
        current_date = datetime.strptime(selected_date, "%d.%m.%Y").date()

        total_parcels = self.parcels_weights_listbox.size()
        total_cost = sum(float(self.parcels_weights_listbox.get(i).split()[0]) for i in range(total_parcels))
        total_cost_with_vat = total_cost * self.NDS  # Учет НДС (20%)

        result_string = (
            f"Итого за {current_date.strftime('%d.%m.%Y')} отправлено посылок: {total_parcels}  на общую сумму "
            f"{total_cost:.2f} руб. (без НДС) и {total_cost_with_vat:.2f} руб. (с НДС).\n"
        )

        custom_path = self.custom_path
        filename = os.path.join(custom_path, f"Списки посылок.txt")

        with open(filename, "a", encoding='utf-8') as file:
            file.write(result_string)

        messagebox.showinfo("Успешно", "Результаты сохранены.")
        self.parcels_window.destroy()

# Функция для создания диалогового окна по общему подсчету и ввода даты
    def ask_month_input(self):
        if hasattr(self, 'month_window') and self.month_window.winfo_exists():
            self.month_window.focus_set()
            return

        self.month_window = tk.Toplevel(self)
        self.month_window.title("Выберите дату для отчета")
        self.month_window.geometry("300x320")

        background_label = tk.Label(self.month_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.month_window.winfo_screenwidth()
        screen_height = self.month_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 300) // 2 + 500
        y_coordinate = (screen_height - 320) // 2

        # Устанавливаем положение окна по центру
        self.month_window.geometry(f"300x320+{x_coordinate}+{y_coordinate}")

        label_text = "День обязателен, но не учитывается"
        label = tk.Label(self.month_window, text=label_text)
        label.pack(pady=(10, 0))
        
        self.month_calendar = Calendar(self.month_window, selectmode='day', year=2024, month=2, date_pattern='dd.mm.yyyy', locale='ru_RU')

        self.month_calendar.pack(pady=10)

        self.ok_button = tk.Button(self.month_window, text="Сохранить", command=self.get_selected_month, bg='lightgray')
        self.ok_button.pack(pady=10, padx=20, fill=tk.X)
        self.ok_button.configure(borderwidth=2, relief=tk.GROOVE)

        self.month_window.geometry(f"300x320+{x_coordinate}+{y_coordinate}")

    def get_selected_month(self):
        selected_date_str = self.month_calendar.get_date()
        
        if selected_date_str:
            day, month, year = map(int, selected_date_str.split('.'))
            selected_month = f"{month:02d}.{year}"
            self.calculate_total_for_month(selected_month)
            self.month_window.destroy()
        else:
            messagebox.showwarning("Внимание", "Дата не выбрана.")

# Это сохранение итогов (за определенный месяц)
    def calculate_total_for_month(self, selected_month):
        try:
            custom_path = self.custom_path
            month, year = map(int, selected_month.split("."))
        
            # Инициализация переменных для подсчета итогов
            total_weight = 0
            total_cost = 0.0
            total_parcels = 0
            total_letters_cost = 0.0
            total_registered_letters_cost = 0.0
            total_simple_letters = 0
            total_registered_letters = 0
            total_parcels_package = 0
            total_cost_package = 0
            total_cost_with_vat_package = 0
        
            # Обработка данных из файла бандеролей
            parcels_file_path = os.path.join(custom_path, "Списки бандеролей.txt")
            if not os.path.exists(parcels_file_path):
                raise FileNotFoundError(f"Файл {parcels_file_path} не найден.")
        
            with open(parcels_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    date_match = re.search(r'(\d{2})\.(\d{2})\.(\d{4})', line)
                    if date_match:
                        line_day, line_month, line_year = map(int, date_match.groups())
                        if line_month == month and line_year == year:
                            total_parcels += int(re.search(r'отправлено (\d+) (бандероль|бандероли|бандеролей)', line).group(1))
                            total_weight += float(re.search(r'весом ([\d.]+) грамм', line).group(1))
                            total_cost += float(re.search(r'на сумму ([\d.]+) (рублей|руб\.)', line).group(1))
        
            # Обработка данных из файла простых писем
            simple_letters_file_path = os.path.join(custom_path, "Списки простых писем.txt")
            if not os.path.exists(simple_letters_file_path):
                raise FileNotFoundError(f"Файл {simple_letters_file_path} не найден.")
        
            with open(simple_letters_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        total_simple_letters += int(re.search(r'Количество писем: (\d+)', line).group(1))
                        total_letters_cost += float(re.search(r'Итого: ([\d.]+) руб.', line).group(1))
        
            # Обработка данных из файла заказных писем
            registered_letters_file_path = os.path.join(custom_path, "Списки заказных писем.txt")
            if not os.path.exists(registered_letters_file_path):
                raise FileNotFoundError(f"Файл {registered_letters_file_path} не найден.")
        
            with open(registered_letters_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        total_registered_letters += int(re.search(r'Количество писем: (\d+)', line).group(1))
                        total_registered_letters_cost += float(re.search(r'Итого: ([\d.]+) руб.', line).group(1))
            
            # Обработка иностранных писем
            foreign_letters_file_path = os.path.join(custom_path, "Списки иностранных писем.txt")
            if os.path.exists(foreign_letters_file_path):
                with open(foreign_letters_file_path, 'r', encoding='utf-8') as foreign_file:
                    foreign_data = foreign_file.readlines()

                # Отфильтруем строки, соответствующие выбранному месяцу и году
                foreign_data_for_month = [line.strip() for line in foreign_data if f"{month:02d}.{year}" in line]

                # Строка для добавления в отчет
                foreign_letters_string = "\n".join(foreign_data_for_month)

            # Обработка посылок
            parcels_file_path = os.path.join(custom_path, "Списки посылок.txt")
            if not os.path.exists(parcels_file_path):
                raise FileNotFoundError(f"Файл {parcels_file_path} не найден.")

            with open(parcels_file_path, 'r', encoding='utf-8') as file:
                for line in file:
                    if selected_month in line:
                        match_parcels = re.findall(r'отправлено посылок: (\d+)|отправлено (\d+) посылок', line)
                        for match in match_parcels:
                            num_sent = match[0] or match[1]
                            total_parcels_package += int(num_sent)
                        match_cost = re.search(r'на общую сумму ([\d.]+) (рублей|руб\.) \(без НДС\)', line)
                        if match_cost:
                            total_cost_package += float(match_cost.group(1))
                        match_cost_with_vat = re.search(r'и ([\d.]+) (рублей|руб\.) \(с НДС\)', line)
                        if match_cost_with_vat:
                            total_cost_with_vat_package += float(match_cost_with_vat.group(1))


            total_cost_package = round(total_cost_package, 2)
            total_cost_with_vat_package = round(total_cost_with_vat_package, 2)

            # Строка результата расчета
            result_string = (f"Итого за {selected_month}:\n"
                             f"Отправлено бандеролей: {total_parcels} "
                             f"весом {total_weight} грамм на сумму {total_cost} руб.\n"
                             f"Отправлено простых писем: {total_simple_letters} на сумму {total_letters_cost} руб.\n"
                             f"Отправлено заказных писем: {total_registered_letters} на сумму {total_registered_letters_cost} руб.\n"
                             f"Отправлено иностранных писем:\n{foreign_letters_string}\n"
                             f"Отправлено посылок: {total_parcels_package} на сумму {total_cost_package} руб. (без НДС) и {total_cost_with_vat_package} руб. (с НДС)\n")
        
            # Показываем результат в message box
            messagebox.showinfo("Итоги за месяц", result_string)
        
            # Сохранение в файл
            output_file_path = os.path.join(custom_path, f"Отчет за {selected_month}.txt")
            with open(output_file_path, 'w', encoding='utf-8') as output_file:
                output_file.write(result_string)
    
        except ValueError as ve:
            messagebox.showerror("Ошибка", str(ve))
        except FileNotFoundError as fe:
            messagebox.showerror("Ошибка", str(fe))

# Кнопка создания обложки
    def open_cover_window(self):
        if hasattr(self, 'cover_window') and self.cover_window.winfo_exists():
            self.cover_window.focus_set()
            return
        self.cover_window = tk.Toplevel(self)
        self.cover_window.title("Выберите тип обложки")

        # Размер окна
        window_width = 270
        window_height = 390

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.cover_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.cover_window.winfo_screenwidth()
        screen_height = self.cover_window.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - window_width) // 2 + 450
        y = (screen_height - window_height) // 2

        # Устанавливаем позицию окна
        self.cover_window.geometry(f"{window_width}x{window_height}+{x}+{y}")

        # Кнопки
        self.parcel_button = tk.Button(self.cover_window, text="      Обложка почты       ", command=self.open_post)
        self.parcel_button.pack(pady=(50, 0))  # Устанавливаем отступ сверху для первой кнопки
        self.parcel_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.parcel_button.configure(bg='lightgrey')

        tk.Label(self.cover_window, text="").pack()

        self.mail_button = tk.Button(self.cover_window, text="    Обложка посылок    ", command=self.open_pacage)
        self.mail_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.mail_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.mail_button.configure(bg='lightgrey')

        tk.Label(self.cover_window, text="").pack()

        self.doc_button = tk.Button(self.cover_window, text=" Обложка документов ", command=self.open_documents)
        self.doc_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.doc_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.doc_button.configure(bg='lightgrey')

        tk.Label(self.cover_window, text="").pack()

        self.complaints_button = tk.Button(self.cover_window, text="      Обложка жалоб      ", command=self.open_complaints)
        self.complaints_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.complaints_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.complaints_button.configure(bg='lightgrey')

        tk.Label(self.cover_window, text="").pack()

        self.invoices_button = tk.Button(self.cover_window, text="  Обложка накладных  ", command=self.open_invoice)
        self.invoices_button.pack(pady=(10, 0),)  # Устанавливаем отступ снизу для второй кнопки
        self.invoices_button.configure(borderwidth=2, relief=tk.GROOVE)
        self.invoices_button.configure(bg='lightgrey')

# Обложка на почту
    def open_post(self):
        if hasattr(self, 'top') and self.top.winfo_exists():
            self.top.focus_set()
            return
        # Создаем окно
        self.top = tk.Toplevel(self)
        self.top.title("Выбор даты")
        
        background_label = tk.Label(self.top, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - self.top.winfo_reqwidth()) / 2
        y = (screen_height - self.top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        self.top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            self.top.destroy()
            self.create_document(selected_date)

        # Календарь для выбора даты
        cal = Calendar(self.top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(self.top, text="Сформировать обложку", command=get_date, bg='lightgray')
        btn_ok.pack(pady=10, padx=20, fill=tk.X)
        btn_ok.configure(borderwidth=2, relief=tk.GROOVE)

    def create_document(self, selected_date):
        # Создаем документ Word с выбранной датой
        document = DocxDocument()

        # Устанавливаем размер страницы A4
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
        
        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", "Реестры","передачи", "почтовых", "отправлений", f"{selected_date.strftime('%d.%m.%Y')}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(70)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        # Устанавливаем ориентацию страницы на книжную
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path)

# Обложка на посылки
    def open_pacage(self):
        if hasattr(self, 'top') and self.top.winfo_exists():
            self.top.focus_set()
            return
        # Создаем окно
        self.top = tk.Toplevel(self)
        self.top.title("Выбор даты")
        
        background_label = tk.Label(self.top, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - self.top.winfo_reqwidth()) / 2
        y = (screen_height - self.top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        self.top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            self.top.destroy()
            self.create_document_2(selected_date)

        # Календарь для выбора даты
        cal = Calendar(self.top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(self.top, text="Сформировать обложку", command=get_date, bg='lightgray')
        btn_ok.pack(pady=10, padx=20, fill=tk.X)
        btn_ok.configure(borderwidth=2, relief=tk.GROOVE)

    def create_document_2(self, selected_date):
        # Создаем документ Word с выбранной датой
        document = DocxDocument()

        # Устанавливаем размер страницы A4
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
    
        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        #Преобразуем номер месяца в текстовый формат
        month_name = calendar.month_name[selected_date.month]
        month_name_ru = {
            'January': 'Январь',
            'February': 'Февраль',
            'March': 'Март',
            'April': 'Апрель',
            'May': 'Май',
            'June': 'Июнь',
            'July': 'Июль',
            'August': 'Август',
            'September': 'Сентябрь',
            'October': 'Октябрь',
            'November': 'Ноябрь',
            'December': 'Декабрь'
        }

        month_name = month_name_ru.get(month_name, month_name)

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", " ", " ", " ", "Посылки", f"{month_name} {selected_date.year}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(72)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        section = document.sections[0]
        new_width, new_height = section.page_height, section.page_width
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width = new_width
        section.page_height = new_height

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path) 

    # определяем первый рабочий день
    def get_first_workday(self, year, month):
        num_days = calendar.monthrange(year, month)[1]  # количество дней в месяце
        for day in range(1, num_days + 1):
            date = datetime(year, month, day)
            if date.weekday() < 5:  # Понедельник - пятница (0-4)
                return date

    # Определяем последний рабочий день
    def get_last_workday(self, year, month):
        num_days = calendar.monthrange(year, month)[1]  # количество дней в месяце
        for day in range(num_days, 0, -1):
            date = datetime(year, month, day)
            if date.weekday() < 5:
                return date

    # Обложка на документы
    def open_documents(self):
        if hasattr(self, 'top') and self.top.winfo_exists():
            self.top.focus_set()
            return
        # Создаем окно
        self.top = tk.Toplevel(self)
        self.top.title("Выбор периода")
        
        background_label = tk.Label(self.top, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.top.winfo_screenwidth()
        screen_height = self.top.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - self.top.winfo_reqwidth()) / 2
        y = (screen_height - self.top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        self.top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            self.top.destroy()
            self.create_document_3(selected_date)

        # Календарь для выбора даты
        cal = Calendar(self.top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(self.top, text="Сформировать обложку", command=get_date, bg='lightgray')
        btn_ok.pack(pady=10, padx=20, fill=tk.X)
        btn_ok.configure(borderwidth=2, relief=tk.GROOVE)

    def create_document_3(self, selected_date):
        # Определяем первый и последний рабочий день в выбранном месяце
        first_workday = self.get_first_workday(selected_date.year, selected_date.month)
        last_workday = self.get_last_workday(selected_date.year, selected_date.month)

        # Преобразуем даты в нужный формат (дд.мм.гггг)
        first_workday_str = first_workday.strftime('%d.%m.%Y')
        last_workday_str = last_workday.strftime('%d.%m.%Y')

        # Создаем документ Word с выбранными датами
        document = DocxDocument()

        # Устанавливаем размер страницы A4
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)
    
        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'
    
        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", "Реестры","передачи", "документов", " ", f"{first_workday_str}-{last_workday_str}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(72)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman


        # Устанавливаем ориентацию страницы на книжную
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path)

    # Обложка на апелляционные жалобы
    def open_complaints(self):
        if hasattr(self, 'top') and self.top.winfo_exists():
            self.top.focus_set()
            return
        # Создаем окно
        self.top = tk.Toplevel(self)
        self.top.title("Выбор периода")
        
        background_label = tk.Label(self.top, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - self.top.winfo_reqwidth()) / 2
        y = (screen_height - self.top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        self.top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            self.top.destroy()
            self.create_document_4(selected_date)

        # Календарь для выбора даты
        cal = Calendar(self.top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(self.top, text="Сформировать обложку", command=get_date, bg='lightgray')
        btn_ok.pack(pady=10, padx=20, fill=tk.X)
        btn_ok.configure(borderwidth=2, relief=tk.GROOVE)

    def create_document_4(self, selected_date):
        # Определяем первый и последний рабочий день в выбранном месяце
        first_workday = self.get_first_workday(selected_date.year, selected_date.month)
        last_workday = self.get_last_workday(selected_date.year, selected_date.month)

        # Преобразуем даты в нужный формат (дд.мм.гггг)
        first_workday_str = first_workday.strftime('%d.%m.%Y')
        last_workday_str = last_workday.strftime('%d.%m.%Y')

        # Создаем документ Word с выбранными датами
        document = DocxDocument()

        # Устанавливаем размер страницы A4
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)

        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", "Реестры","передачи", "а/ж", " ", f"{first_workday_str}-{last_workday_str}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(70)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        # Устанавливаем ориентацию страницы на книжную
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path)

# Обложка на накладные
    def open_invoice(self):
        if hasattr(self, 'top') and self.top.winfo_exists():
            self.top.focus_set()
            return
        # Создаем окно
        self.top = tk.Toplevel(self)
        self.top.title("Выбор периода")
        
        background_label = tk.Label(self.top, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        # Получаем размеры экрана
        screen_width = self.winfo_screenwidth()
        screen_height = self.winfo_screenheight()

        # Вычисляем координаты для отображения окна посередине экрана
        x = (screen_width - self.top.winfo_reqwidth()) / 2
        y = (screen_height - self.top.winfo_reqheight()) / 2

        # Устанавливаем позицию окна
        self.top.geometry("+%d+%d" % (x, y))

        # Функция для выбора даты
        def get_date():
            selected_date_str = cal.get_date()
            selected_date = datetime.strptime(selected_date_str, "%d.%m.%Y")
            self.top.destroy()
            self.create_document_5(selected_date)

        # Календарь для выбора даты
        cal = Calendar(self.top, selectmode="day", year=datetime.now().year, month=datetime.now().month, day=datetime.now().day, locale='ru_RU')
        cal.pack(padx=10, pady=10)

        # Кнопка для выбора даты
        btn_ok = tk.Button(self.top, text="Сформировать обложку", command=get_date, bg='lightgray')
        btn_ok.pack(pady=10, padx=20, fill=tk.X)
        btn_ok.configure(borderwidth=2, relief=tk.GROOVE)

    def create_document_5(self, selected_date):
        # Определяем первый и последний рабочий день в выбранном месяце
        first_workday = self.get_first_workday(selected_date.year, selected_date.month)
        last_workday = self.get_last_workday(selected_date.year, selected_date.month)

        # Преобразуем даты в нужный формат (дд.мм.гггг)
        first_workday_str = first_workday.strftime('%d.%m.%Y')
        last_workday_str = last_workday.strftime('%d.%m.%Y')

        # Создаем документ Word с выбранными датами
        document = DocxDocument()

        # Устанавливаем размер страницы A4
        section = document.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21)

        # Установка шрифта Times New Roman
        run_font = document.styles['Normal'].font
        run_font.name = 'Times New Roman'

        # Добавляем каждую надпись на отдельной строке с выравниванием по центру
        for text in [" ", " ", "Почтовые","накладные", " ", f"{first_workday_str}-{last_workday_str}"]:
            paragraph = document.add_paragraph()
            run = paragraph.add_run(text)
            run.font.size = Pt(72)
            run.font.bold = True  # Установка жирного шрифта
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.space_after = Inches(0.2)  # Пространство после каждой строки
            run.font.name = 'Times New Roman'  # Установка шрифта Times New Roman

        # Устанавливаем ориентацию страницы на книжную
        section = document.sections[0]
        section.orientation = WD_ORIENTATION.PORTRAIT

        # Открываем документ
        temp_file_path = tempfile.mktemp(suffix='.docx')
        document.save(temp_file_path)
        os.startfile(temp_file_path)

# Открывается кнопка Настроек
    def open_settings_window(self):
        if hasattr(self, 'settings_window') and self.settings_window.winfo_exists():
            self.settings_window.focus_set()
            return
        self.settings_window = tk.Toplevel(self)
        self.settings_window.title("Настройки")
        self.settings_window.geometry("400x710")
        self.settings_window.attributes('-topmost', 'true')

        self.developer_email = "law.borzzz@gmail.com"

        # Создаем метку с изображением в качестве фона
        background_label = tk.Label(self.settings_window, image=self.bg_image_tk)
        background_label.place(x=0, y=0, relwidth=1, relheight=1)

        screen_width = self.settings_window.winfo_screenwidth()
        screen_height = self.settings_window.winfo_screenheight()

        # Рассчитываем координаты для центрирования окна
        x_coordinate = (screen_width - 400) // 2 - 500
        y_coordinate = (screen_height - 710) // 2

        # Устанавливаем положение окна по центру
        self.settings_window.geometry(f"400x710+{x_coordinate}+{y_coordinate}")

        self.settings_title_label = tk.Label(self.settings_window, text="Настройки почтовых тарифов")
        self.settings_title_label.pack(pady=(10, 0))  # небольшой отступ сверху

        # Линия для отделения надписи от остальных элементов
        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(5, 10)) 
        
        self.settings_window.attributes('-topmost', 'true')
        
        self.base_cost_label = tk.Label(self.settings_window, text="Стоимость бандероли в 120 грамм:")
        self.base_cost_label.pack(pady=10)
        
        self.base_cost_entry = tk.Entry(self.settings_window)
        self.base_cost_entry.pack(pady=10)
        self.base_cost_entry.insert(0, str(self.BASE_COST))
        
        self.step_cost_label = tk.Label(self.settings_window, text="Стоимость за шаг в 20 грамм:")
        self.step_cost_label.pack(pady=10)
        
        self.step_cost_entry = tk.Entry(self.settings_window)
        self.step_cost_entry.pack(pady=10)
        self.step_cost_entry.insert(0, str(self.STEP_COST))
        
        self.letter_cost_label = tk.Label(self.settings_window, text="Стоимость простого письма:")
        self.letter_cost_label.pack(pady=10)
        
        self.letter_cost_entry = tk.Entry(self.settings_window)
        self.letter_cost_entry.pack(pady=10)
        self.letter_cost_entry.insert(0, str(self.LETTER_COST))

        self.registered_letter_cost_label = tk.Label(self.settings_window, text="Стоимость заказного письма:")
        self.registered_letter_cost_label.pack(pady=10)
        
        self.registered_letter_cost_entry = tk.Entry(self.settings_window)
        self.registered_letter_cost_entry.pack(pady=10)
        self.registered_letter_cost_entry.insert(0, str(self.REGISTERED_LETTER_COST))

        self.nds_entry_label = tk.Label(self.settings_window, text="Переводим надбавку НДС:")
        self.nds_entry_label.pack(pady=10)

        self.nds_entry = tk.Entry(self.settings_window)
        self.nds_entry.pack(pady=10)
        self.nds_entry.insert(0, str(self.NDS))

        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(10, 10))

        self.settings_title_label = tk.Label(self.settings_window, text="Настройка пути хранения файлов.\n Без необходимости не трогать!!!")
        self.settings_title_label.pack(pady=(10, 0))  # небольшой отступ сверху

        self.custom_path_entry = tk.Entry(self.settings_window)
        self.custom_path_entry.pack(pady=10)
        self.custom_path_entry.insert(0, str(self.custom_path))

        self.custom_path_button = tk.Button(self.settings_window, text="Выбрать путь", command=self.select_custom_path, bg='lightgray')
        self.custom_path_button.pack(pady=10, padx=20, fill=tk.X)
        self.custom_path_button.configure(borderwidth=2, relief=tk.GROOVE)

        self.settings_separator = tk.Frame(self.settings_window, height=2, bg="grey")  # создаем рамку с высотой 2 пикселя и цветом серого
        self.settings_separator.pack(fill='x', pady=(10, 10))
        
        self.save_settings_button = tk.Button(
            self.settings_window, text="Сохранить", command=self.save_settings, bg='lightgray'
        )
        self.save_settings_button.pack(pady=10, padx=20, fill=tk.X)
        self.save_settings_button.configure(borderwidth=2, relief=tk.GROOVE)

        # Добавляем метку "О разработчике"
        self.developer_label = tk.Label(self.settings_window, text="О программе", fg="blue", cursor="hand2")
        self.developer_label.pack(side=tk.BOTTOM, anchor=tk.SE, padx=10, pady=10)
        self.developer_label.bind("<Button-1>", lambda event: self.open_program_info())

    def open_program_info(self):
        if hasattr(self, 'program_info_window') and self.program_info_window.winfo_exists():
            self.program_info_window.focus_set()
            return
        self.program_info_window = tk.Toplevel(self)
        self.program_info_window.title("Информация о программе")
        self.program_info_window.geometry("350x250+{}+{}".format(
            (self.winfo_screenwidth() - 350) // 2,
            (self.winfo_screenheight() - 250) // 2
        ))
        
        # Создаем метку с отступом сверху
        top_spacing = tk.Label(self.program_info_window, text="", font=("Helvetica", 2))
        top_spacing.pack()

        # Создаем метку с заголовком отчета почты
        report_label = tk.Label(self.program_info_window, text="Формирование отчёта", font=("Calibri", 16, "bold"))
        report_label.pack(pady=(10, 5))

        # Версия программы
        version_label = tk.Label(self.program_info_window, text="Версия: 4.0.3", font=("Calibri", 12))
        version_label.pack(pady=5)

        # Информация о разработчиках
        developer_label = tk.Label(self.program_info_window, text="Главный разработчик: Борзиков Д.А.", font=("Calibri", 10), anchor="w")
        developer_label.pack(anchor="w", padx=10)

        leader_label = tk.Label(self.program_info_window, text="Руководитель: Усова Н.Н.", font=("Calibri", 10), anchor="w")
        leader_label.pack(anchor="w", padx=10)

        tester_label = tk.Label(self.program_info_window, text="Тестировщик: Паринов М.И.", font=("Calibri", 10), anchor="w")
        tester_label.pack(anchor="w", padx=10)

        # Для связи (адрес электронной почты)
        contact_label = tk.Label(self.program_info_window, text="Почта для связи", font=("Calibri", 10, "bold"), cursor="hand2", fg="blue")
        contact_label.pack(pady=5)

        # Привязываем событие щелчка мыши к метке с адресом электронной почты
        contact_label.bind("<Button-1>", lambda event: self.copy_email_to_clipboard())

        # Дополнительные сведения
        additional_info_label = tk.Label(self.program_info_window, text="Нажмите, чтобы скопировать.", font=("Calibri", 8, "italic"), fg="gray")
        additional_info_label.pack(pady=5)

    def copy_email_to_clipboard(self):
        self.clipboard_clear()  # Очищаем буфер обмена
        self.clipboard_append(self.developer_email)  # Копируем адрес электронной почты
        messagebox.showinfo("Скопировано", "Адрес электронной почты скопирован.")


# Окно выбора пути сохранения
    def select_custom_path(self):
        path = askdirectory()  # Показать диалоговое окно и вернуть выбранный путь
        if path:
            self.custom_path_entry.delete(0, tk.END)  # Очистка текущего содержимого Entry
            self.custom_path_entry.insert(0, path)  # Вставить выбранный путь

# Сохраняем настройки
    def save_settings(self):
        try:
            self.BASE_COST = float(self.base_cost_entry.get())
            self.STEP_COST = float(self.step_cost_entry.get())
            self.LETTER_COST = float(self.letter_cost_entry.get())
            self.NDS = float(self.nds_entry.get())
            self.custom_path = self.custom_path_entry.get()
            self.save_settings_to_file()  # вызов метода для сохранения всех настроек, включая путь
            self.settings_window.destroy()
            messagebox.showinfo("Успех", "Настройки сохранены.")
        except ValueError:
            messagebox.showerror("Ошибка", "Введите корректные числовые значения.")

# Это сохранение настроек, чтобы они не сбивались при закрытии    
    def save_settings_to_file(self):
        settings = {
            'BASE_COST': self.BASE_COST,
            'STEP_COST': self.STEP_COST,
            'LETTER_COST': self.LETTER_COST,
            'REGISTERED_LETTER_COST' : self.REGISTERED_LETTER_COST,
            'NDS' : self.NDS,
            'CUSTOM_PATH': self.custom_path
        }
        settings_path = os.path.join(os.getenv('APPDATA'), 'settings.json')
        os.makedirs(os.path.dirname(settings_path), exist_ok=True)  # Создаем директорию, если она не существует
        with open(settings_path, 'w') as f:
            json.dump(settings, f)

        print(f'Файл настроек сохранен по пути: {settings_path}')

# Отсюда загружаем настройки приложения
    def load_settings_from_file(self):
        settings_path = os.path.join(os.getenv('APPDATA'), 'settings.json')
        try:
            with open(settings_path, 'r') as f:
                settings = json.load(f)
            self.BASE_COST = settings.get('BASE_COST', self.BASE_COST)
            self.STEP_COST = settings.get('STEP_COST', self.STEP_COST)
            self.LETTER_COST = settings.get('LETTER_COST', self.LETTER_COST)
            self.REGISTERED_LETTER_COST = settings.get('REGISTERED_LETTER_COST', self.REGISTERED_LETTER_COST)
            self.NDS = settings.get('NDS', self.NDS)
            self.custom_path = settings.get('CUSTOM_PATH', os.path.expanduser('~'))
        except FileNotFoundError:
            self.custom_path = os.path.expanduser('~') # Файл с настройками отсутствует, будут использованы значения по умолчанию
            self.save_settings_to_file()  # Создаем файл настроек со значениями по умолчанию
        except json.JSONDecodeError:
            messagebox.showerror("Ошибка", "Ошибка чтения настроек. Проверьте файл настроек.")



def main():
    current_version = "4.0.3"  # Текущая версия вашего приложения
    check_for_updates(current_version)
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()